import os
import glob
import logging
import openpyxl
from docx import Document
from PIL import ImageGrab
from docx.shared import Cm
from datetime import datetime
from calendar import monthrange


def pasta_existente(pasta):
    try:
        if not os.path.exists(pasta):
            os.makedirs(pasta)
    except OSError as e:
        print(f"Erro ao criar a pasta {pasta}: {e}")


def fazer_planilha(path_dados_excel, data_atual, quantidade_dias_mes, minha_conexao_download, minha_conexao_upload, speedtest_download, speedtest_upload):

    mes_ano = {
        "1": "Janerio",
        "2": "Fevereiro",
        "3": "Março",
        "4": "Abril",
        "5": "Maio",
        "6": "Junho",
        "7": "Julho",
        "8": "Agosto",
        "9": "Setembro",
        "10": "Outubro",
        "11": "Novembro",
        "12": "Dezembro"
    }
    nome_mes = mes_ano[str(data_atual.month)]
    try:
        # Verifica se o arquivo existe, caso não exista, cria um novo arquivo
        if not os.path.isfile(path_dados_excel):
            planilha = openpyxl.Workbook()
            planilha.save(path_dados_excel)
            planilha.close()

        # Abre o arquivo excel
        planilha = openpyxl.load_workbook(path_dados_excel)
        # Verifica se a planilha do mês e ano atual existe, caso não exista, cria uma nova planilha
        if not f"{nome_mes}_{data_atual.year}" in planilha.sheetnames:
            sheet = planilha.create_sheet(title=f"{nome_mes}_{data_atual.year}")
            sheet["A2"], sheet["B1"], sheet["D1"], sheet["F1"] = ("Dia", "Minha Conexão", "SpeedTest", "Média")
            sheet.merge_cells("B1:C1")
            sheet.merge_cells("D1:E1")
            sheet.merge_cells("F1:G1")
            # Criação das colunas
            for x in range(8):
                if x > 1 and x % 2 == 0:
                    sheet.cell(row=2, column=x).value = "Download"
                elif x > 1 and x % 2 == 1:
                    sheet.cell(row=2, column=x).value = "Upload"
        
        # Seleciona a planilha do mês e ano atual
        planilha.active = planilha[f"{nome_mes}_{data_atual.year}"]
        sheet = planilha.active
        # Preenche a planilha com os dados
        for linha in range(int(quantidade_dias_mes)):
            # Coluna Dia
            if sheet.cell(row=linha + 3, column=1).value is None:
                sheet.cell(row=linha + 3, column=1).value = linha + 1

            if data_atual.day == sheet.cell(row=linha + 3, column=1).value:
                # Coluna Minha Conexao Download
                if sheet.cell(row=linha + 3, column=2).value is None:
                    sheet.cell(row=linha + 3, column=2).value = minha_conexao_download.replace('.', ',')
                # Coluna Minha Conexao Upload
                if sheet.cell(row=linha + 3, column=3).value is None:
                    sheet.cell(row=linha + 3, column=3).value = minha_conexao_upload.replace('.', ',')
                # Coluna Speed Test Download
                if sheet.cell(row=linha + 3, column=4).value is None:
                    sheet.cell(row=linha + 3, column=4).value = speedtest_download.replace('.', ',')
                # Coluna Speed Test Upload
                if sheet.cell(row=linha + 3, column=5).value is None:
                    sheet.cell(row=linha + 3, column=5).value = speedtest_upload.replace('.', ',')
                # Coluna Média
                celula_download_minha_conexao = f"{sheet.cell(row=linha + 3, column=2).coordinate}"
                celula_download_speed_teste = f"{sheet.cell(row=linha + 3, column=4).coordinate}"
                celula_upload_minha_conexao = f"{sheet.cell(row=linha + 3, column=3).coordinate}"
                celula_upload_speed_teste = f"{sheet.cell(row=linha + 3, column=5).coordinate}"
                if sheet.cell(row=linha + 3, column=6).value is None:
                    sheet.cell(row=linha + 3, column=6).value = f'=({celula_download_minha_conexao}+{celula_download_speed_teste})/2'
                if sheet.cell(row=linha + 3, column=7).value is None:
                    sheet.cell(row=linha + 3, column=7).value = f'=({celula_upload_minha_conexao}+{celula_upload_speed_teste})/2'
        
        # Calculo de media mensal
        media_mensal = sheet.cell(row=int(quantidade_dias_mes) + 3, column=1).value = "Velocidade Média em Mbps"
        
        # Calculo de media diaria
        primeira_celula_download = f"{sheet.cell(row=3, column=6).coordinate}"
        ultima_celula_download = f"{sheet.cell(row=int(quantidade_dias_mes) + 2, column=6).coordinate}"

        primeira_celula_upload = f"{sheet.cell(row=3, column=7).coordinate}"
        ultima_celula_upload = f"{sheet.cell(row=int(quantidade_dias_mes) + 2, column=7).coordinate}"
        
        sheet.cell(row=int(quantidade_dias_mes) + 3, column=6).value = f"=SUM({primeira_celula_download}:{ultima_celula_download})/{int(quantidade_dias_mes)}".upper()
        sheet.cell(row=int(quantidade_dias_mes) + 3, column=7).value = f"=SUM({primeira_celula_upload}:{ultima_celula_upload})/{int(quantidade_dias_mes)}".upper()

        celula_media_mensal_download = f"{sheet.cell(row=int(quantidade_dias_mes) + 3, column=6).coordinate}"
        celula_media_mensal_upload = f"{sheet.cell(row=int(quantidade_dias_mes) + 3, column=7).coordinate}"

        # Calculo de media em porcentagem
        sheet.cell(row=int(quantidade_dias_mes) + 4, column=1).value = "Velocidade média em %"
        sheet.cell(row=int(quantidade_dias_mes) + 4,
                column=6).value = f"=({celula_media_mensal_download}*100)/{int(quantidade_dias_mes)}"
        sheet.cell(row=int(quantidade_dias_mes) + 4,
                column=7).value = f"=({celula_media_mensal_upload}*100)/{int(quantidade_dias_mes)}"
        # Salva a planilha
        planilha.save(path_dados_excel)
        planilha.close()
    except Exception as e:
        print("Ocorreu um erro ao tentar fazer a planilha.", "ERRO: ", e)


def fazer_documento_word(path_dados_word, path_prints, data_atual, quantidade_dias_mes):
    try:
        # Verifica se o arquivo existe, caso não exista, cria um novo arquivo
        if not os.path.isfile(path_dados_word):
            documento = Document()
            documento.save(path_dados_word)

        # Abre o arquivo word
        documento = Document(path_dados_word)
        # Adiciona o cabeçalho somente na primeira página
        if len(documento.paragraphs) == 0:
            documento.add_heading("MEDIÇÕES DE VELOCIDADE")
        # Adiciona os prints da tela
        #Entender o que isso faz!
        documento.add_paragraph(f"Data: {data_atual.day}/{data_atual.month}/{data_atual.year}")

        for caminho_print in glob.glob(os.path.abspath(os.path.join(path_prints, "*.png"))):
            documento.add_picture(caminho_print, width=Cm(17.05), height=Cm(10.71))
        documento.save(path_dados_word)
    except Exception as e:
        print("Ocorreu um erro ao tentar criar o arquivo world.", "ERRO: ", e)


def configurar_logger(nome_logger='app', arquivo_log='app.log', nivel=logging.INFO):
    """
        Configura e retorna um logger para gerenciar logs da aplicação.

        Args:
            nome_logger (str): Nome do logger, usado para identificação.
            arquivo_log (str): Caminho do arquivo onde os logs serão salvos.
            nivel (int): Nível de registro (DEBUG, INFO, WARNING, ERROR, CRITICAL).

        Returns:
            logging.Logger: Objeto logger configurado.
    """
    # Criação do diretório para logs, se necessário.
    diretorio_log = os.path.dirname(arquivo_log)
    if diretorio_log and not os.path.exists(diretorio_log):
        os.makedirs(diretorio_log)

    # Configurar o logger
    logger = logging.getLogger(nome_logger)
    logger.setLevel(nivel)

    #Evitar Múltiplos handlers no logger
    if not logger.handlers:
        #formato do logger
        formato = logging.Formatter(
            "%(asctime)s - %(name)s - %(levelname)s - %(message)s"
        )

        #handler para arquivo
        arquivo_handler = logging.FileHandler(arquivo_log)
        arquivo_handler.setFormatter(formato)
        logger.addHandler(arquivo_handler)

        #Handler para console
        console_handler = logging.StreamHandler()
        console_handler.setFormatter(formato)
        logger.addHandler(console_handler)

    return logger