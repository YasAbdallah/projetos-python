import os

import openpyxl.workbook
import openpyxl
import glob
from docx import Document
from docx.shared import Cm
from lib.navegacao import Navegar
from time import sleep
from datetime import datetime
from PIL import ImageGrab
from calendar import monthrange


data_atual = datetime.now()
# monthrange retorna o último dia do mês, basta setá-lo na data e pronto
ultimo_dia = str(data_atual.replace(day=monthrange(data_atual.year, data_atual.month)[1])).split("-")[-1]
quantidade_dias_mes = str(ultimo_dia).split(" ")[0]
mes_ano = {
    "1": "Janerio",
    "2": "Fevereiro",
    "3": "Março",
    "4": "Abril",
    "5": "Maio",
    "6": "Junho",
    "7": "Julho",
    "8": "Agosto",
    "9": "Setembro",
    "10": "Outubro",
    "11": "Novembro",
    "12": "Dezembro"}

CAMINHO_USUARIO = f'D:\\Temp\\internet'
PASTA_DATA_ATUAL = f'{data_atual.year}\\{data_atual.month}\\{data_atual.day}'

if not os.path.exists(os.path.join(CAMINHO_USUARIO, PASTA_DATA_ATUAL)):
    os.makedirs(os.path.join(CAMINHO_USUARIO, PASTA_DATA_ATUAL))
    
driver = Navegar("D:\\scripts\\webDriver")
try:
    driver.abrirSite('https://www.speedtest.net/')
    driver.clicar('//button[text() = "Aceitar cookies" or text() = "Continue" or text() = "Continuar"]')
    driver.clicar('//span[@class="start-text"]')
    sleep(45)
    driver.clicar('//button[text() = "Close"]')
   
    speedtest_download = str(driver.capturarTexto('//span[@class="result-data-large number result-data-value download-speed"]'))
    speedtest_upload = str(driver.capturarTexto('//span[@class="result-data-large number result-data-value upload-speed"]'))

    imagem = ImageGrab.grab()
    imagem.save(os.path.join(CAMINHO_USUARIO, PASTA_DATA_ATUAL, 'speedTeste.png'))
except Exception as e:
    print("Ocorreu um erro ao tentar medir a velocidade da internet do site SpeedTest.")
finally:
    driver.fecharNavegador()

try:
    driver.abrirSite('https://www.minhaconexao.com.br/')
    driver.clicar('//button[@type="button" and text()="Iniciar"]')
    sleep(60)
    minha_conexao_download = str(driver.capturarTexto('/html/body/section[1]/div/div[2]/div[2]/div[7]/div/div[1]/div[2]/div/span[3]')).split(" ")[0]
    minha_conexao_upload = str(driver.capturarTexto('/html/body/section[1]/div/div[2]/div[2]/div[7]/div/div[1]/div[3]/div/span[3]')).split(" ")[0]
    imagem = ImageGrab.grab()
    imagem.save(os.path.join(CAMINHO_USUARIO, PASTA_DATA_ATUAL, 'minhaConexao.png'))
except Exception as e:
    print("Ocorreu um erro ao tentar medir a velocidade da internet no site Minha Conexao.")
finally:
    driver.fecharNavegador() 
    
try:   
    CAMINHO_PLANILHA_MEDICAO = 'D:\\Temp\\Controle mensal de velocidade.xlsx'

    if not os.path.isfile(CAMINHO_PLANILHA_MEDICAO):
        planilha = openpyxl.Workbook()
        planilha.save(CAMINHO_PLANILHA_MEDICAO)
        planilha.close()

    planilha = openpyxl.load_workbook(CAMINHO_PLANILHA_MEDICAO)
    if not f"{mes_ano[str(data_atual.month)]}_{data_atual.year}" in planilha.sheetnames:
        sheet = planilha.create_sheet(title=f"{mes_ano[str(data_atual.month)]}_{data_atual.year}")

        sheet["A2"], sheet["B1"], sheet["D1"], sheet["F1"] = ("Dia", "Minha Conexão", "SpeedTest", "Média")
        sheet.merge_cells("B1:C1")
        sheet.merge_cells("D1:E1")
        sheet.merge_cells("F1:G1")
        for x in range(8):
            if x > 1 and x % 2 == 0:
                sheet.cell(row=2, column=x).value = "Download"
            elif x > 1 and x % 2 == 1:
                sheet.cell(row=2, column=x).value = "Upload"
                
    planilha.active = planilha[f"{mes_ano[str(data_atual.month)]}_{data_atual.year}"]
    sheet = planilha.active
    for linha in range(int(quantidade_dias_mes)):

        # Coluna Dia
        sheet.cell(row=linha+3, column=1).value = linha+1 if sheet.cell(row=linha+3, column=1).value == None else None
            
        if data_atual.day == sheet.cell(row=linha+3, column=1).value:
            #Coluna Minha Conexao Download
            sheet.cell(row=linha+3, column=2).value = minha_conexao_download.replace('.', ',') if sheet.cell(row=linha+3, column=2).value == None else None
                
            #Coluna Minha Conexao Upload
            sheet.cell(row=linha+3, column=3).value = minha_conexao_upload.replace('.', ',') if sheet.cell(row=linha+3, column=3).value == None else None
                
            #Coluna Speed Test Download
            sheet.cell(row=linha+3, column=4).value = speedtest_download.replace('.', ',') if sheet.cell(row=linha+3, column=4).value == None else None
                
            #Coluna Speed Test Download
            sheet.cell(row=linha+3, column=5).value = speedtest_upload.replace('.', ',') if sheet.cell(row=linha+3, column=5).value == None else None

            #Coluna Média
            celula_download_minha_conexao = f"{sheet.cell(row=linha+3, column=2).column_letter}{sheet.cell(row=linha+3, column=2).row}"
            celula_download_speed_teste = f"{sheet.cell(row=linha+3, column=4).column_letter}{sheet.cell(row=linha+3, column=4).row}"

            celula_upload_minha_conexao = f"{sheet.cell(row=linha+3, column=3).column_letter}{sheet.cell(row=linha+3, column=3).row}"
            celula_upload_speed_teste = f"{sheet.cell(row=linha+3, column=5).column_letter}{sheet.cell(row=linha+3, column=5).row}"
            if sheet.cell(row=linha+3, column=6).value == None:
                sheet.cell(row=linha+3, column=6).value = f'=({celula_download_minha_conexao}+{celula_download_speed_teste})/2'
            if sheet.cell(row=linha+3, column=7).value == None:
                sheet.cell(row=linha+3, column=7).value = f'=({celula_upload_minha_conexao}+{celula_upload_speed_teste})/2'
    media_mensal = sheet.cell(row=int(quantidade_dias_mes)+3, column=1).value = "Velocidade Média em Mbps"

    primeira_celula_download = f"{sheet.cell(row=3, column=6).column_letter}{sheet.cell(row=3, column=6).row}"
    ultima_celula_download = f"{sheet.cell(row=int(quantidade_dias_mes)+2, column=6).column_letter}{sheet.cell(row=int(quantidade_dias_mes)+2, column=6).row}"

    primeira_celula_upload = f"{sheet.cell(row=3, column=7).column_letter}{sheet.cell(row=3, column=7).row}"
    ultima_celula_upload = f"{sheet.cell(row=int(quantidade_dias_mes)+2, column=7).column_letter}{+sheet.cell(row=int(quantidade_dias_mes)+2, column=7).row}"

    sheet.cell(row=int(quantidade_dias_mes)+3, column=6).value = f"=SUM({primeira_celula_download}:{ultima_celula_download})/{int(quantidade_dias_mes)}".upper()
    sheet.cell(row=int(quantidade_dias_mes)+3, column=7).value = f"=SUM({primeira_celula_upload}:{ultima_celula_upload})/{int(quantidade_dias_mes)}".upper()


    celula_media_mensal_download = f"{sheet.cell(row=int(quantidade_dias_mes)+3, column=6).column_letter}{sheet.cell(row=int(quantidade_dias_mes)+3, column=6).row}"
    celula_media_mensal_upload = f"{sheet.cell(row=int(quantidade_dias_mes)+3, column=7).column_letter}{sheet.cell(row=int(quantidade_dias_mes)+3, column=7).row}"

    # Calculo de media em porcentagem
    sheet.cell(row=int(quantidade_dias_mes)+4, column=1).value = "Velocidade média em %"
    sheet.cell(row=int(quantidade_dias_mes)+4, column=6).value = f"=({celula_media_mensal_download}*100)/{int(quantidade_dias_mes)}"
    sheet.cell(row=int(quantidade_dias_mes)+4, column=7).value = f"=({celula_media_mensal_upload}*100)/{int(quantidade_dias_mes)}"

    planilha.save(CAMINHO_PLANILHA_MEDICAO)
    planilha.close()
except Exception as e:
    print("Ocorreu um erro ao tentar fazer a planilha.", "ERRO: ", e)

try:
    documento = Document()

    documento.add_heading("MEDIÇÕES DE VELOCIDADE")
    for dia in range(int(quantidade_dias_mes)):
        documento.add_paragraph(f'{dia+1}/{data_atual.month}/{data_atual.year}')
        if os.path.exists(f"D:\\Temp\\5\\{dia+1}"):
            for caminho_print in glob.glob(f"D:\\Temp\\5\\{dia+1}\\*.png"):
                documento.add_picture(caminho_print, width=Cm(17.05), height=Cm(10.71))
        documento.add_page_break()
    documento.save(f"D:/Temp/{data_atual.month}-{data_atual.year}(Medicoes de Velocidade).docx")
except Exception as e:
    print("Ocorreu um erro ao tentar criar o arquivo world.", "ERRO: ", e)